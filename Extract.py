import torch
from transformers import LongformerTokenizer, EncoderDecoderModel
import sumy
import nltk
nltk.download('punkt')
from sumy.parsers.plaintext import PlaintextParser
from spacy.tokens import Span
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from transformers import BartTokenizer, BartForConditionalGeneration
from gingerit.gingerit import GingerIt
import textstat
import docx
import re
from  ispassive import Tagger
t = Tagger()



class extractSummary:
    # Methods
    def set_seed(seed):
        torch.manual_seed(seed)
        if torch.cuda.is_available():
            torch.cuda.manual_seed_all(seed)

    def firstSummary(self, docu):
        all_sentences = []
        #doc = docx.Document('Rise of the Bhakti Movement.docx')
        doc = docx.Document(docu)
        text = ""
        for para in doc.paragraphs:
            text = text+para.text
        model = EncoderDecoderModel.from_pretrained(
            "patrickvonplaten/longformer2roberta-cnn_dailymail-fp16")
        tokenizer = LongformerTokenizer.from_pretrained(
            "allenai/longformer-base-4096")
        device = torch.device('cpu')
        input_ids = tokenizer(text, return_tensors="pt").input_ids
        output_ids = model.generate(input_ids)
        summary = tokenizer.decode(output_ids[0], skip_special_tokens=True)
        return summary

    def secondSummary(self,docu):
        all_sentences = []
        #doc = docx.Document('Rise of the Bhakti Movement.docx')
        doc = docx.Document(docu)
        text = ""
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'):             
                text = text+para.text
        parser = PlaintextParser.from_string(text,Tokenizer("english"))
        summarizer = LexRankSummarizer()
        #Summarize the document with 2 sentences
        summary = summarizer(parser.document, 2)
        newSummary=list(summary)
        text=""
        s=""
        newSummary=' '.join(str(s) for s in summarizer(parser.document, 2))
        #newSummary.append(textLength)
        lenSummary=(newSummary.split())
        return newSummary,len(lenSummary)
    def thirdSummary(self, docu):
        #doc = docx.Document('Rise of the Bhakti Movement.docx')
        doc = docx.Document(docu)
        text = ""
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'): 
                text = text+para.text
        model = BartForConditionalGeneration.from_pretrained(
            'facebook/bart-large-cnn')
        tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')

        # tokenize without truncation
        inputs_no_trunc = tokenizer(
            text, max_length=None, return_tensors='pt', truncation=False)

        # get batches of tokens corresponding to the exact model_max_length
        chunk_start = 0
        chunk_end = tokenizer.model_max_length  # == 1024 for Bart
        inputs_batch_lst = []
        while chunk_start <= len(inputs_no_trunc['input_ids'][0]):
            # get batch of n tokens
            inputs_batch = inputs_no_trunc['input_ids'][0][chunk_start:chunk_end]
            inputs_batch = torch.unsqueeze(inputs_batch, 0)
            inputs_batch_lst.append(inputs_batch)
            chunk_start += tokenizer.model_max_length  # == 1024 for Bart
            chunk_end += tokenizer.model_max_length  # == 1024 for Bart

        # generate a summary on each batch
        summary_ids_lst = [model.generate(
            inputs, num_beams=4, max_length=100, early_stopping=True) for inputs in inputs_batch_lst]

        # decode the output and join into one string with one paragraph per summary batch
        summary_batch_lst = []
        for summary_id in summary_ids_lst:
            summary_batch = [tokenizer.decode(
                g, skip_special_tokens=True, clean_up_tokenization_spaces=False) for g in summary_id]
            summary_batch_lst.append(summary_batch[0])
        summary_all = '\n'.join(summary_batch_lst)
        return summary_all

    def extractHeaders(self, docu):
        headerDict = {}
        heading1 = 0
        heading2 = 0
        heading3 = 0
        heading4 = 0
        heading5 = 0
        heading6 = 0
        document = docx.Document(docu)
        def iter_headings(paragraphs):
            for paragraph in paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    yield paragraph

        for heading in iter_headings(document.paragraphs):
            if heading.style.name == 'Heading 1':
                heading1 = heading1+1
            if heading.style.name == 'Heading 2':
                heading2 = heading2+1
            if heading.style.name == 'Heading 3':
                heading3 = heading3+1
            if heading.style.name == 'Heading 4':
                heading4 = heading4+1
            if heading.style.name == 'Heading 5':
                heading5 = heading5+1
            if heading.style.name == 'Heading 6':
                heading6 = heading6+1
        headerDict={'Headers':['Heading1','Heading2','Heading3','Heading4','Heading5'], 
                   'Total Header Types':[heading1,heading2,heading3,heading4,heading5]}        
        # headerDict['Heading1'] = heading1
        # headerDict['Heading2'] = heading2
        # headerDict['Heading3'] = heading3
        # headerDict['Heading4'] = heading4
        # headerDict['Heading5'] = heading5
        # headerDict['Heading5'] = heading6
        return headerDict

    def extractWordLengthGt20(self, docu):
        doc = docx.Document(docu)
        sentenc=[]
        lnth=[]
        sentenceDict = {}
        all_sentences=[]
        for para in doc.paragraphs:
            all_sentences.append(para.text.split('.'))
        i=0    
        for sentenceWords in all_sentences:
            totalWds= sentenceWords[0].split()     
            sentenceLength = len(totalWds)
            if sentenceLength > 20:
                sentenc.append(sentenceWords[0])
                lnth.append(sentenceLength)

                i=i+1
        sentenceDict['Sentences']=sentenc
        sentenceDict['Length of Sentences']=lnth       
        return sentenceDict

    def getTransitionWordDensity(self, docu):
        totalSentences = 0
        noOfTransitionWords = 0
        all_sentences=[]
        doc = docx.Document(docu)
        transitionList = ['ABOVE', 'ABOVE ALL', 'THUS', 'ACCORDINGLY', 'ACROSS', 'ADJACENT', 'AFTER', 'AFTERWARD', 'AFTERWARDS', 'AGAIN', 'ALONG THE EDGE', 'ALSO', 'ALTHOUGH', 'ANALOGOUS TO', 'AND', 'ANOTHER', 'AROUND', 'AS A RESULT', 'AS AN ILLUSTRATION', 'AS WELL AS', 'AT ANY RATE', 'AT FIRST', 'AT LAST', 'AT LEAST', 'AT LENGTH', 'AT THE BOTTOM', 'AT THE FRONT', 'AT THE LEFT', 'AT THE REAR', 'AT THE RIGHT', 'AT THE SAME TIME', 'AT THE TOP', 'BECAUSE', 'BEFORE',
                          'BEHIND', 'BELOW', 'BENEATH', 'BESIDE', 'BESIDES', 'BEYOND', 'BOTH-AND', 'BUT', 'CERTAINLY', 'CONCURRENTLY', 'CONSEQUENTLY', 'CONTRARILY', 'CONTRAST', 'CONVERSELY', 'DAY', 'DUE TO', 'DURING THE MORNING', 'EQUALLY IMPORTANT', 'ESPECIALLY', 'ETC', 'EVEN THOUGH', 'EVENTUALLY', 'FINALLY', 'FIRST', 'FOR A MINUTE', 'FOR EXAMPLE', 'FOR INSTANCE', 'FOR THIS PURPOSE', 'FOR THIS REASON', 'FORMERLY', 'FURTHER', 'FURTHERMORE', 'GENERALLY', 'GRANTED THAT', 'HENCE']
        for para in doc.paragraphs:
            all_sentences.append(para.text.split('.'))
        totalSentences = len(all_sentences)
            #print("Length : ",len(all_sentences))
        for words in all_sentences:
                #print("Length : ",len(cleanWords))
                #print(cleanWords.split(' ',1))
            cleanWords=words[0].split(' ')    
            if cleanWords[0].upper() in transitionList:
                    noOfTransitionWords = noOfTransitionWords+1
        percentageTransition = float(noOfTransitionWords)*(100/totalSentences)
        percentageTransitionWords = '{:.2f}'.format(percentageTransition)
        return percentageTransitionWords

    def consecutiveSentenceDetect(self, docu):
        all_sentences = []
        doc = docx.Document(docu)
        for para in doc.paragraphs:
            all_sentences.append(para.text.split('.'))
            #print("Length : ",len(all_sentences))
        s=set()
        for words in all_sentences:
            
            cleanWords = ""
            firstWord = ""
            secondWord = ""
            thirdWord = ""
            finalDict={}
            # print(cleanWords.join(words).strip())
            index = all_sentences.index(words)
            #print (str(index))
            if index+2 < len(all_sentences):
             if ' '.join(all_sentences[index]).split(' ', 1)[0].upper() == ' '.join(all_sentences[index+1]).split(' ', 1)[0].upper() == ' '.join(all_sentences[index+2]).split(' ', 1)[0].upper():
                   s.add(' '.join(all_sentences[index]))
                   s.add(' '.join(all_sentences[index+1]))
                   s.add(' '.join(all_sentences[index+2]))
        output =list(s)
        output.sort()
        return output           
    def determineHeadingParalength(self, docu):
        dicthead={}
        doc = docx.Document(docu)
        heading=False
        textNormal=[]   
        textHeader=[]
        newtext=""
        for para in doc.paragraphs:
            if not para.style.name.startswith("Heading"):
                  
                  newtext=newtext+" "+para.text
                  
            else:
               if  para.style.name.startswith("Heading"):
                  if len(newtext)!=0:
                     temp=newtext.split()   
                     textNormal.append(len(temp)) 
                     newtext=""
                  textHeader.append(para.text)


        temp=newtext.split()   
        textNormal.append(len(temp))
        dicthead['Headers']=textHeader
        dicthead['Word Length']=textNormal
        output=dicthead
        # output = dict(zip(textHeader,textNormal)) 
        return output

    def getFleschScore(self, docu):
        easeOfReading={}
        all_sentences=[]
        doc = docx.Document(docu)
        text1 = ""
        for para in doc.paragraphs:
            if not para.style.name.startswith("Heading"):
                all_sentences.append(para.text)
        finalSentence="".join(all_sentences)
        fk=textstat.flesch_kincaid_grade(finalSentence)
        #fk=r.flesch_kincaid()
        print(fk)   
        if float(fk)>=90:
            easeOfReading['Grade']=["Easy for 5th Grade"]
            easeOfReading['Grade Score']=['{:.2f}'.format(float(fk))+ " %"]
        if float(fk)>=80 and float(fk)<90:
            easeOfReading['Grade']=["Easy for 6th Grade"]
            easeOfReading['Grade Score']=['{:.2f}'.format(float(fk))+ " %"]
        if float(fk)>=70 and float(fk)<80:
            easeOfReading['Grade']=["Easy for 7th Grade"]
            easeOfReading['Grade Score']=['{:.2f}'.format(float(fk))+ " %"]
        if float(fk)>=60 and float(fk)<70:
            easeOfReading['Grade']=["Easy for 8th and 9th Grade"]
            easeOfReading['Grade Score']=['{:.2f}'.format(float(fk))+ " %"]
        if float(fk)>=50 and float(fk)<60:
            easeOfReading['Grade']=["Easy for 10th to 12th Grade"]
            easeOfReading['Grade Score']=['{:.2f}'.format(float(fk))+ " %"]
        if float(fk)>=30 and float(fk)<50:
            easeOfReading['Grade']=["Easy for College Grade Students"]
            easeOfReading['Grade Score']=['{:.2f}'.format(float(fk))+ " %"]
        if float(fk)>=0 and float(fk)<30:
            easeOfReading['Grade']=["Easy for only High Intellectuals"]
            easeOfReading['Grade Score']=['{:.2f}'.format(float(fk))+ " %"]
         
        return easeOfReading

    def identifyPassiveForm(self,docu):
        # style = [0=Casual to Formal, 1=Formal to Casual, 2=Active to Passive, 3=Passive to Active etc..]
        doc = docx.Document(docu)
        dictAnalysis={}
        dictStatements={}
        passive=[]
        total=[]
        temp=[]
        for para in doc.paragraphs:
          total.append(para.text)
          temp=para.text.split(".")
          
          if len(temp)>2:
              return "Error..Please Indent document, 1 line per paragraph",temp
          if  t.is_passive(para.text):
              passive.append(para.text)
        
        percentagePassiveWords = float(len(passive))*(100/len(total))
        percentagePassive = '{:.2f}'.format(percentagePassiveWords)
        dictAnalysis['Headers']=['Passive Percentage','Total Passive','Total Sentences']
        dictAnalysis['Header Values']=[percentagePassive+'%',str(len(passive)),str(len(total))]
        dictStatements['Passive Sentences List']=passive
        
        return dictAnalysis, dictStatements

    def identifyImages(self,docu):
        paraGr = []             
        index = []
        doc = docx.Document(docu)
        para=doc.paragraphs
        for i in range(len(doc.paragraphs)):
             if 'graphicData' in para[i]._p.xml:
                 index.append(i)                   
        return len(index)