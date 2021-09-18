import numpy as np
import streamlit as st
import time
import streamlit as st
import streamlit.components.v1 as components
import pandas as  pd
import docx
import os
import Extract
import ExtractFirstSummary
import docx
@st.cache(suppress_st_warning=True)
def color_cell2(df):
   return 'background-color: red'

def color_cell3(df):
   return 'background-color: white'

def color_cell(df):
   if df==0: 
    return 'background-color: red'

def color_cell1(df):
   if df>300: 
    return 'background-color: red'

listOfFunctions=["Setup Summary","Highlight Headers","Sentences Greater than 20 Words","Transition Word Density",
        "Sentences with Same Word Begin","Total Word Length under a Header","Flesch Score",
        "Identify Passive Word Sentences","Identify Total Images in the Document"]

def limitSummaryLength(summ):
    while len(summ)>160:
        opList =summ.split(".")
        lnth=len(opList)   
        opList.pop(lnth-1)
        opList.pop(lnth-2)
        summ= ". ".join(opList)+". "
    return summ


def decideFirstSummary(summaryFirst):
    all_sentences=[]
    doc = docx.Document(summaryFirst)
    for para in doc.paragraphs:
        if not para.style.name.startswith('Heading'): 
            all_sentences.append(para.text)
            sentences = []
    for sentence in all_sentences:
        sentences.append(sentence.replace("[^a-zA-Z]", " ").split(" "))
    totalWords=[]
    for sentence in sentences:
        for first in sentence:
            totalWords.append(first)
    if len(totalWords)<700:
        summ= Extract.extractSummary().secondSummary(summaryFirst)
        return summ

def setupSummary(nameOfDoc):
        twoSummaryDict={}
        twoSummaryDict['Headers']=['First Summary','Second Summary']
        placeholder=st.empty() 
        placeholder.header('Generating probable 2 Summaries for Document')
        with st.spinner('Processing Summaries...'):             
             summ= decideFirstSummary(nameOfDoc)
             if summ!=None:
               finalValue=list(summ)
               finalValueString=limitSummaryLength(finalValue[0])
             else:
               finalValueString="Cannot generate First Summary"   
             summ1= Extract.extractSummary().thirdSummary(nameOfDoc)
             summ1=limitSummaryLength(summ1)
             finalValueString1=summ1
             twoSummaryDict['Summary']=[finalValueString,finalValueString1]
             df = pd.DataFrame(twoSummaryDict)
             st.table(df)

def imagesInDoc(nameOfDoc):
        
        placeholder=st.empty() 
        placeholder.header('Images in the Document  (IID): Ideally More Than 1')
        output=Extract.extractSummary().identifyImages(nameOfDoc)
        with st.spinner('Processing Images...'):
             if (output<=1):
               sentences= "You have less than 2 Images in the Blog"  
               st.write("<style>orange{color:orange}</style>", unsafe_allow_html=True)
               colors=['orange']
               for color in colors:
                    st.write(f"<{color}> {sentences} </{color}>", unsafe_allow_html=True)
             else:       
               sentences= "You have "+str(output) + " images in the Blog"  
               st.write("<style>green{color:green}</style>", unsafe_allow_html=True)
               colors=['green']
               for color in colors:
                    st.write(f"<{color}> {sentences} </{color}>", unsafe_allow_html=True)


def passiveWordSentence(nameOfDoc):
      placeholder=st.empty() 
      placeholder.header('Passive Voice Sentences  (PVS) Ideal PVS 10%')
      dictionary, dictionary1= Extract.extractSummary().identifyPassiveForm(nameOfDoc)
      if dictionary=="Error..Please Indent document, 1 line per paragraph":
         print(dictionary)
         st.write(dictionary)
         st.write(dictionary1)
      else:   
         with st.spinner('Passive Voice Sentences  (PVS)...'):
             df = pd.DataFrame(dictionary)
             temp=dictionary['Header Values'][0]
             temp1=f"{temp[0: -1]}"
             if float(temp1)>10:
                st.table(df.style.applymap(color_cell2, subset=(0,'Header Values')).applymap(color_cell, subset=(0,'Headers')))
             else:
                st.table(df.style.applymap(color_cell3, subset=(0,'Header Values')).applymap(color_cell, subset=(0,'Headers')))
             df = pd.DataFrame(dictionary1)
             st.table(df)

def fleschScore(nameOfDoc):
        placeholder=st.empty() 
        placeholder.header('FLESCH Grading  (FLESCH)')
        dictionary= Extract.extractSummary().getFleschScore(nameOfDoc)
        with st.spinner('Processing FLESCH Grades...'):
             df = pd.DataFrame(dictionary)
             st.table(df)
def wordLengthUnderHeader(nameOfDoc):
        placeholder=st.empty() 
        placeholder.header('Words Under a Heading (WUH) [UNDER 300]')
        dictionary= Extract.extractSummary().determineHeadingParalength(nameOfDoc)
        with st.spinner('Processing Words Under Headers...'):
             dictionary= Extract.extractSummary().determineHeadingParalength(nameOfDoc)
             df = pd.DataFrame(dictionary)
            #  st.table(df)
             st.table(df.style.applymap(color_cell1, subset=('Word Length')))

def sameWordBegin(nameOfDoc):
        placeholder=st.empty() 
        placeholder.header('Consecutive 3 sentences With Same Word Begin (SWB)')
        output=Extract.extractSummary().consecutiveSentenceDetect(nameOfDoc)
        for sentences in output:
            st.write("<style>orange{color:orange}</style>", unsafe_allow_html=True)
            colors=['orange']
            for color in colors:
                    st.write(f"<{color}> {sentences} </{color}>", unsafe_allow_html=True)
   
def transitionWD(nameOfDoc):
        placeholder=st.empty() 
        placeholder.header('Transition Word Density (TWD)')
        with st.spinner('Processing Transition Word Density...'):
              percentTrans=Extract.extractSummary().getTransitionWordDensity(nameOfDoc)
              if (float(percentTrans)< 2 or float(percentTrans)>20):
                   st.write("<style>orange{color:orange}</style>", unsafe_allow_html=True)
                   colors=['orange']
                   for color in colors:
                    st.write(f"<{color}>Ideal TWD should be between (2 and 20%) : {percentTrans} %</{color}>", unsafe_allow_html=True)
              else:
                   st.write("<style>white{color:white}</style>", unsafe_allow_html=True)
                   colors=['white']
                   for color in colors:
                       st.write(f"<{color}>Ideal TWD should be between (2 and 20%) : {percentTrans} %</{color}>", unsafe_allow_html=True)

def sentenceGt20Wd(nameOfDoc):
        placeholder=st.empty() 
        placeholder.header('Sentences Greater than 20 Words (WO20)')
        with st.spinner('Processing Sentences...'):
             dictionary= Extract.extractSummary().extractWordLengthGt20(nameOfDoc)
             df = pd.DataFrame(dictionary)
             st.table(df)    

def highLightHeaders(nameOfDoc):
    placeholder=st.empty() 
    placeholder.header('Headers in the Document(HID)')
    placeholder1=st.empty()   
    with st.spinner('Processing Headers...'):
         dictionary= Extract.extractSummary().extractHeaders(nameOfDoc)
         df = pd.DataFrame(dictionary)
         st.table(df.style.applymap(color_cell, subset=(0,'Total Header Types')).applymap(color_cell, subset=(0,'Headers')))

def getListofDocx():
    listDocs.append('-----Select Document-----')
    for docList in os.listdir('.'):
        if docList.endswith('.docx') :
            listDocs.append(docList)
    return listDocs

listDocs=[]
functionList=[]        
dictListOfFunctions={}
listDocs=getListofDocx()
dictListOfFunctions['List of Documents']=listDocs
st.header('Document To Process')
docOption = st.selectbox('',
        dictListOfFunctions['List of Documents'])
for listOfOptions in listOfFunctions:
     if  st.checkbox(listOfOptions):
         functionList.append(listOfOptions)
for opts in functionList:
    if opts=="Highlight Headers":
       highLightHeaders(docOption)

    if opts=="Sentences Greater than 20 Words":
       sentenceGt20Wd(docOption)

    if opts=="Transition Word Density":
       transitionWD(docOption)

    if opts=="Sentences with Same Word Begin":
       sameWordBegin(docOption)

    if opts=="Total Word Length under a Header":
       wordLengthUnderHeader(docOption)

    if opts=="Flesch Score":
       fleschScore(docOption)

    if opts=="Identify Passive Word Sentences":
        passiveWordSentence(docOption)
    if opts=="Identify Total Images in the Document":
        imagesInDoc(docOption)
    if opts=="Setup Summary":
         setupSummary(docOption)


# def processDocument():
#     with use_scope('scope1', clear=True):
#          for nameOfFunction in pin['ThisCheckbox']:
#              if nameOfFunction=="Setup Summary":
#                 setupSummary(pin['selectDocument'])
#              if nameOfFunction=="Total Word Length under a Header":
#                 wordLengthUnderHeader(pin['selectDocument'])
#              if nameOfFunction=="Flesch Score":
#                 fleschScore(pin['selectDocument'])
#              if nameOfFunction=="Identify Passive Word Sentences":
#                 passiveWordSentence(pin['selectDocument'])
#                 # if nameOfFunction=="Identify Total Images in the Document":
#                 #     imagesInDoc(pin['selectDocument'])

# def limitSummaryLength(summ):
#     #print("This ONE SUMM : ",summ)  
#     while len(summ)>160:
#         opList =summ.split(".")
#         lnth=len(opList)   
#         opList.pop(lnth-1)
#         opList.pop(lnth-2)
#         summ= ". ".join(opList)+". "
#         #print(summ,len(summ))
#     return summ

# def decideFirstSummary(summaryFirst):
#     all_sentences=[]
#     doc = docx.Document(summaryFirst)
#     for para in doc.paragraphs:
#         if not para.style.name.startswith('Heading'): 
#             all_sentences.append(para.text)
#             sentences = []
#     for sentence in all_sentences:
#         sentences.append(sentence.replace("[^a-zA-Z]", " ").split(" "))
#     totalWords=[]
#     for sentence in sentences:
#         for first in sentence:
#             totalWords.append(first)
#     if len(totalWords)<700:
#         summ= Extract.extractSummary().secondSummary(summaryFirst)
#         #print("This OP: 1 List", len("".join(list(summ[0]))))
        
#         if len("".join(list(summ[0])))>160:
#             summ=limitSummaryLength("".join(list(summ[0]))) 
#         else:
#             summ ="".join(list(summ[0]))      
#         return summ
#     else:
#         summ0= ExtractFirstSummary.generate_summary(summaryFirst,4)
#         #print("This OP: 2", summ0)
#         if len("".join(list(summ0[0])))>160:
#             summ0=limitSummaryLength("".join(list(summ0[0])))
#         else:
#             summ0= "".join(list(summ0[0]))         
#         return summ0
# def getFirstSummary():
#         decideFirstSummary(pin['selectDocument'])
#         summ= Extract.extractSummary().secondSummary(pin['selectDocument'])
#         finalValue=list(summ)
#         finalValueString=str(finalValue[1])
#         #print(finalValue[0])
# finalValue=[]    
# finalValueString=''
# finalValueString1=''
# listDocs=getListofDocx()
# # put_markdown('## Select A Document')
# # put_select(name='selectDocument',options=listDocs)
# # put_markdown('## Get One Summary')
# # put_textarea('md_text', rows=20)
# # put_buttons(['Analyze Document'], lambda _: put_text(pin.pin_name))
# put_markdown('## Select A Document')
# put_select(name='selectDocument',options=listDocs)
# put_checkbox(name="ThisCheckbox",options=listOfFunctions, value=listOfFunctions, inline=False)
# put_buttons(['Process Document'], onclick=[processDocument])
# def setupSummary(nameOfDoc):
#         with use_scope('scope11'):
#               put_text('Processing for First Summary (FS). . .')
#               put_loading(shape='grow', color='primary')
#         summ= decideFirstSummary(nameOfDoc)
#         with use_scope('scope11', clear=True):  # enter the existing scope and clear the previous content
#              put_text('Processing for Second Summary (SS). . .')
#              put_loading(shape='grow', color='secondary')
#         finalValue=summ
#         finalValueString=str(finalValue)
#         summ1= Extract.extractSummary().thirdSummary(nameOfDoc)
#         #finalValue1=list(summ1)
#         summ1=limitSummaryLength(summ1)
#         finalValueString1=summ1
#         with use_scope('scope11', clear=True):  # enter the existing scope and clear the previous content
#              put_text('Here is the Result . . .')
#         put_grid([
#           [put_markdown('## Get First Summary (FS)'),put_markdown('## Get Second Summary (SS)')],
#           [put_text(finalValueString), put_text(finalValueString1)]], cell_width='500px').show

# def wordLengthUnderHeader(nameOfDoc):
#         put_markdown('## Words Under a Heading (WUH)')
#         with use_scope('scope15'):
#               put_text('Processing Words Under a Heading (WUH). . .')
#               put_loading(shape='grow', color='primary')
#         dictionary= Extract.extractSummary().determineHeadingParalength(nameOfDoc)
#         with use_scope('scope15', clear=True):
#              put_text('Here are the Number of Words Under the Headings (WUH). . .')
#              put_grid([
#                 [style(put_markdown('### Headers'),'color:green'),style(put_markdown('### No of Words'),'color:blue')]],cell_width='300px').show
    
#              for key, value in dictionary.items():
#                  if int(value)<300:        
#                     put_grid([
#                        [style(put_text(key),'color:green'),style(put_text(value),'color:blue')]],cell_width='300px').show
#                  else:
#                     put_grid([
#                        [style(put_text(key),'color:green'),style(put_text(value),'color:red')]],cell_width='300px').show

# def sentenceGt20Wd(nameOfDoc):
#         put_markdown('## Sentences Greater than 20 Words (WO20)')
#         with use_scope('scope16'):
#               put_text('Processing Words Over 20 in a Sentence (WO20) . . .')
#               put_loading(shape='grow', color='primary')
#         dictionary= Extract.extractSummary().extractWordLengthGt20(nameOfDoc)
#         with use_scope('scope16', clear=True):
#              put_text('Here are the Words Over 20 in a Sentence (WO20) . . .')
#              put_grid([
#                 [style(put_markdown('### Sentences'),'color:blue'),style(put_markdown('### No of Words'),'color:red')]],cell_width='400px').show
    
#              for key, value in dictionary.items():
#                     put_grid([
#                        [style(put_text(value[0]),'color:blue'),style(put_text(value[1]),'color:red')]],cell_width='400px').show

# def fleschScore(nameOfDoc):
#         put_markdown('## FLESCH Grading  (FLESCH)')
#         with use_scope('scope17'):
#               put_text('Processing FLESCH Grading  (FLESCH) . . .')
#               put_loading(shape='grow', color='primary')
#         dictionary= Extract.extractSummary().getFleschScore(nameOfDoc)
#         with use_scope('scope17', clear=True):
#              put_text('Here is FLESCH Grading  (FLESCH) . . .')
#              put_grid([
#                 [style(put_markdown('### Ease of reading'),'color:blue'),style(put_markdown('### Flesch Score'),'color:red')]],cell_width='400px').show
    
#              for key, value in dictionary.items():
#                     put_grid([
#                        [style(put_text(value[0]),'color:blue'),style(put_text(value[1]),'color:red')]],cell_width='400px').show
# def passiveWordSentence(nameOfDoc):
#         put_markdown('## Identify Passive Voice Sentences  (PVS)')
#         with use_scope('scope18'):
#               put_text('Processing for Passive Voice  (PVS) . . .')
#               put_loading(shape='grow', color='primary')
#         dictionary= Extract.extractSummary().identifyPassiveForm(nameOfDoc)
#         with use_scope('scope18', clear=True):
#              for pp in dictionary["Percentage Passive"]:
#                  pp=pp
#              for tp in dictionary["Total Passive"]:
#                  tp=tp
#              for t in dictionary["Total Sentences"]:     
#                  t=t
#              put_text('Here are Passive Voice Sentences (PVS) . . .')
#              put_markdown('## Passive Sentences percent should be below 21%')
#              put_grid([
#                 [style(put_markdown('### Percentage Passive'),'color:blue'),style(put_markdown('### Total Passive'),'color:blue'),style(put_markdown('### Total'),'color:blue')]],cell_width='400px').show
                
#              put_grid([
#                 [style(put_text(pp),'color:green'),style(put_text(tp),'color:green'),style(put_text(t),'color:green')]],cell_width='400px').show
#              for item in dictionary['Passive Sentences List']:
#                     new="\n\n".join(item)
#                     put_grid([
#                        [style(put_text(new),'color:blue')]],cell_width='400px').show

# def transitionWD(nameOfDoc):
#         put_markdown('## Transition Word Density (TWD)')
#         with use_scope('scope13'):
#               put_text('Processing the (TWD) Information . . .')
#               put_loading(shape='grow', color='primary')
#               percentTrans=Extract.extractSummary().getTransitionWordDensity(nameOfDoc)
#         with use_scope('scope13', clear=True):
#               put_text('Here is the (TWD) Information . . .')    
#               if (float(percentTrans)< 2 or float(percentTrans)>20):
#                   style(put_markdown('### Ideal TWD should be between 2 and 20% : '+str(percentTrans)+" %"),'color:red')
#               else:
#                   style(put_markdown('### Ideal TWD should be between 2 and 20% : '+str(percentTrans)+" %"),'color:green')    

# def sameWordBegin(nameOfDoc):
#         put_markdown('## Consecutive sentences With Same Word Begin (SWB)')
#         with use_scope('scope14'):
#              put_text('Processing the (SWB) Information . . .')
#              put_loading(shape='grow', color='primary')
#              output=Extract.extractSummary().consecutiveSentenceDetect(nameOfDoc)

#         with use_scope('scope14', clear=True):
#              put_text('Here is the (SWB) Information . . .')    
#              for sentences in output:
#                   style(put_markdown('#### '+ sentences),'color:red')

# students = {'One': ['jack', '34', 'Sydeny'] ,
#             'Two': ['Riti', '30', 'Delhi' ] ,
#              'Three':['Aadi', '16', 'New York'] }
# dfObj = pd.DataFrame(students)
# st.table(dfObj)
# # embed streamlit docs in a streamlit app
# components.iframe("https://docs.streamlit.io/en/latest")
# name="Rajesh"
# form = st.form(key='my_form')
# form.text_input(label='Enter some text')
# submit_button = form.form_submit_button(label='Submit') 
# if submit_button:
#     st.write(f'hello {name}')    
# progress_bar = st.progress(0)
# status_text = st.empty()
# chart = st.line_chart(np.random.randn(10, 2))

# for i in range(100):
#     # Update progress bar.
#     progress_bar.progress(i + 1)

#     new_rows = np.random.randn(10, 2)

#     # Update status text.
#     status_text.text(
#         'The latest random number is: %s' % new_rows[-1, 1])

#     # Append data to the chart.
#     chart.add_rows(new_rows)

#     # Pretend we're doing some computation that takes time.
#     time.sleep(0.1)

# status_text.text('Done!')
